# -*- coding: utf-8 -*-
"""
### For Google Colab Jupyter Notebook
# Copyrighted by Rventima

# El programa necesita la plantilla y los datos, y cambiar el nombre de cada archivo en las primeras variables. Deben estar en la raíz de este archivo.
### Los marcadores en la plantilla deben tener el siguiente formato <<marcador>>
### y las columndas de la cabecera del .csv deben tener exactamente el mismo nombre que "marcador".
"""
"""
# The program needs the template and data, and the name of each file must be changed in the first variables. They must be in the root of this file.
### The markers in the template must follow the format: <<marker>>
### and the header columns of the .csv file must have exactly the same name as "marker".
"""

nombre_doc_datos = "ejemplo.csv"
nombre_doc_plantilla = "plantilla.docx"

import pandas as pd
!apt-get install libreoffice-writer
import subprocess
!pip install python-docx
from docx import Document
import os

dir_datos = "/content/" + nombre_doc_datos
dir_plantilla = "/content/" + nombre_doc_plantilla
dir_pdf = "/content/salidas_pdf"
dir_docx = "/content/salidas_docx"

os.makedirs(dir_pdf, exist_ok=True)
os.makedirs(dir_docx, exist_ok=True)

doc = Document(dir_plantilla)  # Carga el documento
datos = pd.read_csv(dir_datos)  # Carga los datos

def reemplaza_marcador(documento, nombre_marcador, texto_reemplazo, marcador_especial = False):
  formato_marcador = "<<" + nombre_marcador + ">>"
  for table in documento.tables:
    for row in table.rows:
      for cell in row.cells:
        for paragraph in cell.paragraphs:
          if marcador_especial:
            paragraph.text = paragraph.text.replace(formato_marcador, str(texto_reemplazo))
          else:
            for run in paragraph.runs:
              if formato_marcador in run.text:
                run.text = run.text.replace(formato_marcador, str(texto_reemplazo))

def mostrar_doc(doc):
 for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        print(cell.text)

def extraer_fecha_evaluacion(fecha):
  return fecha.split("/")

import re
def limpiar_documento(documento):
  patron_marcador = r'<<(.*?)>>'
  for table in documento.tables:
    for row in table.rows:
      for cell in row.cells:
        for paragraph in cell.paragraphs:
          coincidencia = re.findall(patron_marcador, paragraph.text)
          if coincidencia:
            for marcador in coincidencia:
              reemplaza_marcador(documento, marcador, " ")

def genera_doc(marcadores, datos : list, index):

  lista_marcadores_actividades = datos.columns[9:12]

  for row, data in datos.iterrows():
    if row == index:
      documento_copia = Document(dir_plantilla)
      for marcador in marcadores:

          if marcador in lista_marcadores_actividades:
            nuevo_marcador = marcador + str(data[marcador])
            reemplaza_marcador(documento_copia, nuevo_marcador, "X", True)

          if marcador == "Grado":
            nuevo_marcador = marcador + str(data[marcador])
            reemplaza_marcador(documento_copia, nuevo_marcador, "X", True)

          if marcador == "Recomendacion":
            nuevo_marcador = marcador + str(data[marcador])
            reemplaza_marcador(documento_copia, nuevo_marcador, "X", True)

          if marcador == "FechaEvaluacion":
            numeros_fecha = extraer_fecha_evaluacion(data[marcador])
            reemplaza_marcador(documento_copia, marcador + "DD", numeros_fecha[0])
            reemplaza_marcador(documento_copia, marcador + "MM", numeros_fecha[1])
            reemplaza_marcador(documento_copia, marcador + "AA", numeros_fecha[2])

          else:
            reemplaza_marcador(documento_copia, marcador, data[marcador])

  limpiar_documento(documento_copia)
  mostrar_doc(documento_copia)
  return documento_copia

def generar_nombre_doc(index):
  apellido_paterno = datos.iloc[index]["ApellidoPaterno"].split(" ")
  apellido_materno = datos.iloc[index]["ApellidoMaterno"].split(" ")
  nombre = datos.iloc[index]["Nombre"].split(" ")

  nombre_completo = ""
  for token in apellido_paterno:
    nombre_completo += token + " "
  for token in apellido_materno:
    nombre_completo += token + " "
  for token in nombre:
    nombre_completo += token + " "

  return nombre_completo.replace(" ", "")

def convertir_a_pdf(input_dir, output_dir):
    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir",
        output_dir,
        input_dir
    ])

lista_marcadores = datos.columns.tolist()

for index in range(len(datos)):

  doc = genera_doc(lista_marcadores, datos, index)

  doc.save(f"{dir_docx}/{generar_nombre_doc(index)}.docx")
  convertir_a_pdf(f"{dir_docx}/{generar_nombre_doc(index)}.docx", dir_pdf)